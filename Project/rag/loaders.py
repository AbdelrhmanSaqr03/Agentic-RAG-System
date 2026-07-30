"""
Document loading layer.
Responsible for detecting supported files in the documents folder and
extracting raw text content from each supported file type.
"""

import os
from typing import List, Dict

import fitz  # PyMuPDF
import docx
import pandas as pd

from config import settings
from utils.helpers import clean_text, get_file_extension
from utils.logger import get_logger

logger = get_logger(__name__)


class DocumentLoader:
    """Loads and extracts text from documents of multiple supported formats."""

    def __init__(self, documents_dir: str = None):
        self.documents_dir = documents_dir or settings.paths.documents_dir

    def list_supported_files(self) -> List[str]:
        """
        Scan the documents directory and return paths of supported files.

        Returns:
            A list of absolute file paths for all supported documents.
        """
        if not os.path.isdir(self.documents_dir):
            logger.warning("Documents directory does not exist: %s", self.documents_dir)
            return []

        files = []
        for file_name in os.listdir(self.documents_dir):
            extension = f".{get_file_extension(file_name)}"
            if extension in settings.supported_extensions:
                files.append(os.path.join(self.documents_dir, file_name))
        return sorted(files)

    def load_all(self) -> List[Dict]:
        """
        Load and extract text for every supported document found.

        Returns:
            A list of dicts: {"source": file_name, "text": extracted_text}
        """
        documents = []
        for file_path in self.list_supported_files():
            try:
                text = self._extract_text(file_path)
                if text.strip():
                    documents.append({
                        "source": os.path.basename(file_path),
                        "text": clean_text(text),
                    })
                else:
                    logger.warning("No text extracted from %s", file_path)
            except Exception as exc:
                logger.error("Failed to load %s: %s", file_path, exc)
        return documents

    def _extract_text(self, file_path: str) -> str:
        """Dispatch extraction based on file extension."""
        extension = get_file_extension(file_path)

        if extension == "pdf":
            return self._extract_pdf(file_path)
        if extension == "docx":
            return self._extract_docx(file_path)
        if extension == "txt":
            return self._extract_txt(file_path)
        if extension == "csv":
            return self._extract_csv(file_path)
        if extension == "xlsx":
            return self._extract_xlsx(file_path)

        raise ValueError(f"Unsupported file extension: {extension}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """Extract text from a PDF file using PyMuPDF."""
        text_parts = []
        with fitz.open(file_path) as pdf_document:
            for page_number, page in enumerate(pdf_document, start=1):
                page_text = page.get_text("text")
                text_parts.append(f"[Page {page_number}]\n{page_text}")
        return "\n".join(text_parts)

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extract text from a Word document."""
        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        tables_text = []
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                tables_text.append(row_text)
        return "\n".join(paragraphs + tables_text)

    @staticmethod
    def _extract_txt(file_path: str) -> str:
        """Extract text from a plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
            return file.read()

    @staticmethod
    def _extract_csv(file_path: str) -> str:
        """Extract a textual representation of a CSV file."""
        data_frame = pd.read_csv(file_path)
        summary = (
            f"CSV file with {len(data_frame)} rows and {len(data_frame.columns)} columns.\n"
            f"Columns: {', '.join(str(c) for c in data_frame.columns)}\n\n"
        )
        return summary + data_frame.to_string(index=False, max_rows=500)

    @staticmethod
    def _extract_xlsx(file_path: str) -> str:
        """Extract a textual representation of every sheet in an Excel file."""
        sheets = pd.read_excel(file_path, sheet_name=None, engine="openpyxl")
        text_parts = []
        for sheet_name, data_frame in sheets.items():
            text_parts.append(
                f"Sheet: {sheet_name} ({len(data_frame)} rows, "
                f"{len(data_frame.columns)} columns)\n"
                f"Columns: {', '.join(str(c) for c in data_frame.columns)}\n"
                f"{data_frame.to_string(index=False, max_rows=500)}"
            )
        return "\n\n".join(text_parts)
